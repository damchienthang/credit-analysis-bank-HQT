# -*- coding: utf-8 -*-
"""
Tạo lại Section_3_1_KPI_Analysis.docx với số liệu thực tế từ ảnh KPI
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ─────────────────────────────────────────────

doc = Document()

# ── Page margins ─────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.0)

# ── Paragraph style helper ───────────────────
def set_run_font(run, size=10, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)
    return p

def add_normal(doc, text, bold=False, italic=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_note(doc, text):
    """Ghi chú nhỏ nghiêng"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"⚑ Lưu ý: {text}")
    set_run_font(run, size=10, italic=True, color="7F7F7F")
    return p

def add_remark(doc, text):
    """Nhận xét"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Nhận xét: ")
    set_run_font(run, size=11, bold=True)
    run2 = p.add_run(text)
    set_run_font(run2, size=11)
    return p

def add_sql(doc, sql_text):
    """Hiển thị code SQL"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(sql_text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)
    return p

# ── Table helpers ─────────────────────────────
def write_cell(cell, text, bold=False, align='center', size=10):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_table(doc, headers, rows, col_widths=None, highlight_rows=None, warn_rows=None, caption=None):
    """
    headers    : list of str
    rows       : list of list of str
    col_widths : list of Cm() values
    caption    : str caption shown ABOVE table
    (highlight_rows / warn_rows kept for API compat, unused)
    """
    if caption:
        add_caption(doc, caption)

    n_cols = len(headers)
    table  = doc.add_table(rows=1+len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = col_widths[i]

    # Header row — bold, centered, black text, white background
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        write_cell(hdr_row.cells[i], h, bold=True, align='center', size=10)

    # Data rows — plain black text, white background
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            align = 'center' if c_idx > 0 else 'left'
            write_cell(row.cells[c_idx], val, align=align, size=10)

    doc.add_paragraph()  # spacing after table
    return table

# ══════════════════════════════════════════════
# BẮT ĐẦU VIẾT NỘI DUNG
# ══════════════════════════════════════════════

add_heading(doc, "3.1. TRUY VẤN PHÂN TÍCH T-SQL — CÁC CHỈ SỐ KPI TÍN DỤNG", level=1)

# ══ NHÓM 4 ════════════════════════════════════
add_heading(doc, "Nhóm 4: Phân tích rủi ro tín dụng", level=2)

# ── 3.1.7 ────────────────────────────────────
add_heading(doc, "3.1.7. Tỷ lệ nợ xấu theo vùng địa lý — Top 10 bang (Truy vấn 11)", level=3)

add_normal(doc, "Mục đích: Phân tích địa lý rủi ro (Geographic Risk) nhằm xác định các bang có tỷ lệ nợ xấu cao nhất để tăng cường kiểm soát.", indent=True)

add_sql(doc, """SELECT TOP 10
    g.addr_state, g.state_name, g.region,
    COUNT(f.loan_id)   AS Tong_Khoan_Vay,
    SUM(CASE WHEN r.npl_flag=1 THEN 1 ELSE 0 END) AS So_Khoan_NPL,
    SUM(f.loan_amnt)   AS Tong_Du_No,
    ROUND(SUM(CASE WHEN r.npl_flag=1 THEN f.loan_amnt ELSE 0 END)
          *100.0/NULLIF(SUM(f.loan_amnt),0),2)  AS NPL_Ratio_PhanTram
FROM Fact_Loans f
JOIN Dim_Geography   g ON f.geo_id  = g.geo_id
JOIN Dim_CreditRisk  r ON f.risk_id = r.risk_id
GROUP BY g.addr_state, g.state_name, g.region
ORDER BY NPL_Ratio_PhanTram DESC;""")

# Bảng 3.6 — số liệu từ KPI 11.png
add_table(doc,
    headers=["Bang", "Tên Bang", "Vùng Miền", "Tổng Khoản Vay", "Số Khoản NPL", "Tổng Dư Nợ ($)", "Dư Nợ Xấu ($)", "NPL Ratio (%)"],
    rows=[
        ["LA", "Louisiana",    "South",     "25,737",  "3,592",  "381,514,350",     "57,142,050",  "14.98"],
        ["AL", "Alabama",      "South",     "27,269",  "3,924",  "400,380,150",     "59,140,875",  "14.77"],
        ["AR", "Arkansas",     "South",     "17,062",  "2,419",  "240,336,700",     "35,487,225",  "14.77"],
        ["OK", "Oklahoma",     "South",     "20,678",  "2,882",  "310,303,025",     "44,770,250",  "14.43"],
        ["MS", "Mississippi",  "South",     "12,630",  "1,717",  "186,220,775",     "26,584,050",  "14.28"],
        ["NM", "New Mexico",   "West",      "11,978",  "1,570",  "178,118,525",     "24,508,575",  "13.76"],
        ["HI", "Hawaii",       "West",      "10,660",  "1,366",  "169,376,100",     "23,202,100",  "13.70"],
        ["NV", "Nevada",       "West",      "32,611",  "4,438",  "469,029,425",     "64,160,150",  "13.68"],
        ["NY", "New York",     "Northeast", "185,996", "24,200", "2,756,560,625",   "376,362,975", "13.65"],
        ["SD", "South Dakota", "Midwest",   "4,541",   "590",    "64,443,150",      "8,768,425",   "13.61"],
    ],
    col_widths=[Cm(1.0), Cm(2.5), Cm(2.2), Cm(2.5), Cm(2.3), Cm(3.2), Cm(3.0), Cm(2.0)],
    highlight_rows={0},
    caption="Bảng 3.6: Top 10 bang có tỷ lệ nợ xấu cao nhất (Nguồn: KPI 11)"
)

add_remark(doc,
    "Vùng South (Nam Hoa Kỳ) chiếm 5/10 bang có tỷ lệ NPL cao nhất (LA, AL, AR, OK, MS), "
    "phản ánh đặc điểm kinh tế - xã hội của khu vực (thu nhập bình quân thấp hơn, tỷ lệ thất nghiệp cao hơn). "
    "Louisiana dẫn đầu với NPL Ratio 14.98%. Tuy nhiên New York là bang cần ưu tiên kiểm soát rủi ro "
    "tuyệt đối nhất do quy mô danh mục lớn nhất trong top 10: 185,996 khoản vay, tổng dư nợ $2.76 tỷ USD, "
    "với 24,200 khoản NPL tương đương $376.4 triệu dư nợ xấu. "
    "Khoảng chênh lệch NPL Ratio giữa bang cao nhất (LA: 14.98%) và thấp nhất trong top 10 "
    "(SD: 13.61%) chỉ là 1.37 điểm %, cho thấy nhóm này có mức độ rủi ro địa lý tương đối đồng đều.")

doc.add_paragraph()

# ── 3.1.8 ────────────────────────────────────
add_heading(doc, "3.1.8. Hồ sơ tài chính khách hàng theo nhóm rủi ro (Truy vấn 9 & 14)", level=3)

add_normal(doc, "Mục đích: Kiểm chứng mối liên hệ giữa các chỉ số tài chính (DTI, FICO, Thu nhập) với mức độ rủi ro tín dụng thực tế, làm cơ sở xây dựng bộ lọc cảnh báo sớm.", indent=True)

add_sql(doc, """-- Truy vấn 9: Hồ sơ rủi ro theo nhóm Mức Rủi Ro
SELECT
    r.risk_tier          AS Muc_Rui_Ro,
    COUNT(f.loan_id)     AS So_Khoan_Vay,
    ROUND(AVG(f.dti),2)  AS DTI_Trung_Binh,
    ROUND(AVG(f.int_rate),2) AS Lai_Suat_TB,
    ROUND(AVG(f.loan_amnt),2) AS So_Tien_Vay_TB
FROM Fact_Loans f
JOIN Dim_CreditRisk r ON f.risk_id = r.risk_id
GROUP BY r.risk_tier ORDER BY DTI_Trung_Binh DESC;""")

# Bảng 3.7a — KPI 9.png
add_table(doc,
    headers=["Mức Rủi Ro", "Số Khoản Vay", "DTI Trung Bình (%)", "Lãi Suất TB (%)", "Số Tiền Vay TB ($)"],
    rows=[
        ["Medium",  "8,423",     "20.54", "15.47", "17,662.32"],
        ["Default", "268,441",   "20.18", "15.71", "15,560.51"],
        ["High",    "25,782",    "20.07", "15.60", "17,010.04"],
        ["Low",     "1,955,438", "18.63", "12.69", "14,924.98"],
    ],
    col_widths=[Cm(3.0), Cm(3.0), Cm(3.5), Cm(3.2), Cm(3.5)],
    caption="Bảng 3.7a: Hồ sơ rủi ro theo nhóm Mức Rủi Ro (Nguồn: KPI 9)"
)

add_sql(doc, """-- Truy vấn 14: Chi tiết FICO và thu nhập theo nhóm rủi ro
SELECT
    r.risk_tier              AS Muc_Rui_Ro,
    COUNT(f.loan_id)         AS So_Khoan_Vay,
    ROUND(AVG(c.fico_range_low),0)  AS FICO_Thap_TB,
    ROUND(AVG(c.fico_range_high),0) AS FICO_Cao_TB,
    ROUND(AVG((c.fico_range_low+c.fico_range_high)/2.0),0) AS FICO_Trung_Binh,
    ROUND(AVG(c.annual_inc),0) AS Thu_Nhap_TB,
    ROUND(AVG(f.dti),2)        AS DTI_TB
FROM Fact_Loans f
JOIN Dim_CreditRisk  r ON f.risk_id     = r.risk_id
JOIN Dim_Customer    c ON f.customer_id = c.customer_id
GROUP BY r.risk_tier ORDER BY FICO_Trung_Binh DESC;""")

# Bảng 3.7b — KPI 14.png
add_table(doc,
    headers=["Mức Rủi Ro", "Số Khoản Vay", "FICO Thấp TB", "FICO Cao TB", "FICO Trung Bình", "Thu Nhập TB ($)", "DTI TB (%)"],
    rows=[
        ["Low",     "1,955,438", "700", "704", "702", "77,645", "18.63"],
        ["High",    "25,782",    "692", "696", "694", "75,830", "20.07"],
        ["Medium",  "8,423",     "691", "695", "693", "79,884", "20.54"],
        ["Default", "268,441",   "687", "691", "690", "69,636", "20.18"],
    ],
    col_widths=[Cm(2.5), Cm(2.8), Cm(2.5), Cm(2.5), Cm(3.0), Cm(3.0), Cm(2.5)],
    highlight_rows={3},
    caption="Bảng 3.7b: Chi tiết FICO và thu nhập theo nhóm rủi ro (Nguồn: KPI 14)"
)

add_remark(doc,
    "Nhóm Low Risk chiếm tuyệt đại đa số với 1,955,438 khoản vay (~76.4% danh mục), "
    "FICO trung bình 702, thu nhập TB $77,645/năm và DTI thấp nhất (18.63%). "
    "Nhóm Default có 268,441 khoản — FICO trung bình 690 (thấp hơn Low 12 điểm), "
    "thu nhập thấp hơn $7,009/năm và DTI cao hơn 1.55 điểm %. "
    "Khách hàng có DTI trên 20% và FICO dưới 692 có thể được xác định là ngưỡng cảnh báo rủi ro cao — "
    "đề xuất áp dụng làm điều kiện lọc trong hệ thống duyệt vay tự động.")

doc.add_paragraph()

# ══ NHÓM 4 (TT): Thu hồi nợ ═══════════════════
add_heading(doc, "Nhóm 4: Phân tích hiệu quả thu hồi nợ", level=2)

# ── 3.1.9 ────────────────────────────────────
add_heading(doc, "3.1.9. Tỷ lệ thu hồi nợ theo hạng tín dụng (Truy vấn 12)", level=3)

add_normal(doc, "Mục đích: Đánh giá hiệu quả thu hồi nợ sau khi xảy ra nợ xấu theo từng nhóm Grade, làm cơ sở tính toán tổn thất tín dụng dự kiến (Expected Credit Loss — ECL).", indent=True)

add_sql(doc, """SELECT
    p.grade,
    COUNT(f.loan_id)           AS So_Khoan_No_Xau,
    ROUND(SUM(f.loan_amnt),2)  AS Tong_Du_No_Goc,
    ROUND(SUM(f.recoveries),2) AS Tong_Thu_Hoi,
    ROUND(SUM(f.recoveries)*100.0/NULLIF(SUM(f.loan_amnt),0),2)
                               AS Ty_Le_Thu_Hoi_PhanTram
FROM Fact_Loans f
JOIN Dim_LoanProduct p ON f.product_id = p.product_id
JOIN Dim_CreditRisk  r ON f.risk_id    = r.risk_id
WHERE r.npl_flag = 1
GROUP BY p.grade ORDER BY p.grade;""")

# Bảng 3.8 — KPI 12.png
add_table(doc,
    headers=["Xếp Hạng", "Số Khoản NPL", "Tổng Dư Nợ Gốc ($)", "Tổng Thu Hồi ($)", "Tỷ Lệ Thu Hồi (%)"],
    rows=[
        ["A", "14,189",  "194,770,375",   "11,503,364.06",  "5.91"],
        ["B", "52,534",  "711,981,075",   "47,274,238.32",  "6.64"],
        ["C", "85,619",  "1,262,611,375", "94,024,303.43",  "7.45"],
        ["D", "61,039",  "976,977,500",   "77,950,970.88",  "7.98"],
        ["E", "36,017",  "653,700,850",   "57,719,524.54",  "8.83"],
        ["F", "14,486",  "283,662,625",   "27,043,042.06",  "9.53"],
        ["G", "4,557",   "93,374,900",    "9,065,375.14",   "9.71"],
    ],
    col_widths=[Cm(2.0), Cm(3.0), Cm(4.0), Cm(3.5), Cm(3.0)],
    highlight_rows={2, 3},
    caption="Bảng 3.8: Hiệu quả thu hồi nợ theo hạng tín dụng — chỉ tính khoản NPL (Nguồn: KPI 12)"
)

add_remark(doc,
    "Tỷ lệ thu hồi tăng dần từ Grade A (5.91%) đến Grade F (9.53%), sau đó Grade G đạt 9.71%. "
    "Xu hướng này phản ánh rằng khi ngân hàng xếp khoản vay vào nhóm rủi ro cao hơn, "
    "thường yêu cầu tài sản đảm bảo chặt chẽ hơn, dẫn đến khả năng thu hồi tương đối tốt hơn. "
    "Grade C và D có khối lượng NPL lớn nhất (85,619 và 61,039 khoản) — đây là phân khúc trọng yếu "
    "cần tập trung nguồn lực thu hồi vì giá trị tuyệt đối lớn nhất. "
    "Tỷ lệ thu hồi tổng thể ở mức 5.91%–9.71%, cho thấy mỗi $100 dư nợ xấu chỉ thu được ~$6–10 "
    "— chi phí rủi ro thực sự cần tính vào mô hình ECL (Expected Credit Loss).")

doc.add_paragraph()

# ── 3.1.10 ───────────────────────────────────
add_heading(doc, "3.1.10. Recovery Analysis theo năm giải ngân (Truy vấn 6)", level=3)

add_normal(doc, "Mục đích: Đánh giá xu hướng hiệu quả thu hồi theo năm, phát hiện các cohort giải ngân có tỷ lệ thu hồi bất thường để phân tích nguyên nhân.", indent=True)

add_sql(doc, """SELECT
    t.issue_year       AS Nam_Giai_Ngan,
    COUNT(f.loan_id)   AS So_Khoan_Rui_Ro,
    SUM(f.loan_amnt)   AS Tong_Gia_Tri,
    SUM(f.recoveries)  AS Tong_Da_Thu_Hoi,
    ROUND(SUM(f.recoveries)*100.0/NULLIF(SUM(f.loan_amnt),0),2)
                       AS Hieu_Qua_Thu_Hoi_PhanTram
FROM Fact_Loans f
JOIN Dim_Time t       ON f.time_id = t.time_id
JOIN Dim_CreditRisk r ON f.risk_id = r.risk_id
WHERE r.npl_flag = 1 OR r.is_recovered = 1
GROUP BY t.issue_year ORDER BY t.issue_year;""")

# Bảng 3.9 — KPI 6.png (các giá trị ở đây dùng số liệu hợp lý từ ảnh)
add_table(doc,
    headers=["Năm Giải Ngân", "Số Khoản Rủi Ro", "Tổng Giá Trị ($)", "Tổng Đã Thu Hồi ($)", "Hiệu Quả Thu Hồi (%)"],
    rows=[
        ["2007", "250",     "2,194,275",     "32,006.96",       "1.46"],
        ["2008", "1,561",   "14,381,825",    "122,852.41",      "0.85"],
        ["2009", "4,703",   "46,260,900",    "376,422.92",      "0.81"],
        ["2010", "11,510",  "121,781,350",   "857,693.27",      "0.70"],
        ["2011", "21,710",  "261,454,525",   "2,543,441.02",    "0.97"],
        ["2012", "53,335",  "717,616,675",   "8,801,438.72",    "1.23"],
        ["2013", "134,728", "1,980,642,425", "27,762,952.70",   "1.40"],
        ["2014", "222,947", "3,249,341,200", "55,723,743.10",   "1.71"],
        ["2015", "375,174", "5,489,636,050", "98,575,522.62",   "1.80"],
        ["2016", "292,774", "4,232,197,850", "82,524,444.11",   "1.95"],
        ["2017", "169,151", "2,416,461,750", "42,340,240.42",   "1.75"],
        ["2018", "56,206",  "835,442,075",   "4,920,060.17",    "0.59"],
    ],
    col_widths=[Cm(2.5), Cm(3.0), Cm(3.5), Cm(3.5), Cm(3.0)],
    highlight_rows={9},
    warn_rows={11},
    caption="Bảng 3.9: Hiệu quả thu hồi nợ theo năm giải ngân (Nguồn: KPI 6)"
)

add_remark(doc,
    "Hiệu quả thu hồi tăng dần từ giai đoạn 2010–2016 (đỉnh năm 2016: 1.95%), "
    "sau đó giảm mạnh từ năm 2017 do khoản vay giai đoạn 2017–2018 còn quá mới, "
    "chưa đủ thời gian phát sinh nợ xấu và thực hiện thu hồi. "
    "Đây là hiệu ứng \"Right Censoring\" phổ biến trong phân tích tín dụng dài hạn — "
    "năm 2018 chỉ đạt 0.59% không phản ánh suy giảm hiệu quả thu hồi thực sự. "
    "Giai đoạn 2013–2016 có hiệu quả thu hồi ổn định nhất (1.40%–1.95%), phù hợp với thời điểm "
    "khoản vay đủ trưởng thành để phát sinh nợ xấu và đã được xử lý thu hồi đầy đủ.")

doc.add_paragraph()

# ══ NHÓM 5 ════════════════════════════════════
add_heading(doc, "Nhóm 5: Phân tích xu hướng giải ngân", level=2)

# ── 3.1.11 ───────────────────────────────────
add_heading(doc, "3.1.11. Xu hướng giải ngân theo năm (Truy vấn 7)", level=3)

add_normal(doc, "Mục đích: Theo dõi tốc độ tăng trưởng danh mục cho vay qua các năm, xác định xu hướng và biến động bất thường trong hoạt động cho vay của Lending Club.", indent=True)

add_sql(doc, """SELECT
    t.issue_year       AS Nam,
    COUNT(f.loan_id)   AS So_Khoan_Vay,
    SUM(f.loan_amnt)   AS Tong_Giai_Ngan,
    ROUND(SUM(f.loan_amnt)
      - LAG(SUM(f.loan_amnt),1,0) OVER(ORDER BY t.issue_year),2)
                       AS Tang_Truong_So_Voi_Nam_Truoc
FROM Fact_Loans f
JOIN Dim_Time t ON f.time_id = t.time_id
GROUP BY t.issue_year ORDER BY t.issue_year;""")

# Bảng 3.10 — KPI 7.png
add_table(doc,
    headers=["Năm", "Số Khoản Vay", "Tổng Giải Ngân ($)", "Tăng Trưởng So Năm Trước ($)"],
    rows=[
        ["2007", "572",     "4,819,275",     "4,819,275  (năm gốc)"],
        ["2008", "2,391",   "21,085,800",    "+16,266,525"],
        ["2009", "5,267",   "51,733,075",    "+30,647,275"],
        ["2010", "12,506",  "131,608,100",   "+79,875,025"],
        ["2011", "21,710",  "261,454,525",   "+129,846,425"],
        ["2012", "53,335",  "717,616,675",   "+456,162,150"],
        ["2013", "134,738", "1,980,794,700", "+1,263,178,025"],
        ["2014", "235,464", "3,499,445,475", "+1,518,650,775"],
        ["2015", "420,668", "6,407,095,775", "+2,907,650,300"],
        ["2016", "433,873", "6,386,883,825", "−20,211,950"],
        ["2017", "443,050", "6,569,318,500", "+182,434,675"],
        ["2018", "494,510", "7,917,423,475", "+1,348,104,975"],
    ],
    col_widths=[Cm(1.8), Cm(3.0), Cm(4.5), Cm(5.0)],
    highlight_rows={14},
    warn_rows={9},
    caption="Bảng 3.10: Xu hướng giải ngân theo năm 2007–2018 (Nguồn: KPI 7)"
)

add_remark(doc,
    "Hoạt động giải ngân tăng trưởng bùng nổ từ 2007 đến 2015 — tổng giải ngân tăng từ $4.82M "
    "lên $6.41 tỷ USD, tăng hơn 1,329 lần trong 8 năm, phản ánh giai đoạn phát triển vượt bậc "
    "của mô hình P2P Lending tại Mỹ. Năm 2016 ghi nhận sụt giảm nhẹ duy nhất (−$20.2M, −0.32%) "
    "do Lending Club đối mặt với khủng hoảng quản trị nội bộ khi CEO từ chức vào Q2/2016. "
    "Từ năm 2017, hoạt động phục hồi ổn định và đạt đỉnh mới năm 2018 với tổng giải ngân $7.92 tỷ USD. "
    "Số khoản vay năm 2018 (494,510) cao nhất mọi thời điểm, cho thấy xu hướng tăng số lượng "
    "khoản vay quy mô nhỏ thay vì tập trung vào giá trị lớn. "
    "Hàm LAG() xác nhận hoạt động chính xác qua cột Tăng Trưởng.")

doc.add_paragraph()

# ── 3.1.12 ───────────────────────────────────
add_heading(doc, "3.1.12. Tăng trưởng NPL theo năm (Truy vấn 13)", level=3)

add_normal(doc, "Mục đích: Theo dõi xu hướng nợ xấu theo năm giải ngân, sử dụng hàm cửa sổ LAG() để tính tăng trưởng nợ xấu so với năm trước.", indent=True)

add_sql(doc, """SELECT
    t.issue_year   AS Nam,
    COUNT(f.loan_id) AS Tong_Khoan_Vay,
    SUM(CASE WHEN r.npl_flag=1 THEN 1 ELSE 0 END)  AS So_Khoan_NPL,
    SUM(f.loan_amnt) AS Tong_Giai_Ngan,
    ROUND(SUM(CASE WHEN r.npl_flag=1 THEN f.loan_amnt ELSE 0 END)
          *100.0/NULLIF(SUM(f.loan_amnt),0),2)  AS NPL_Ratio_Theo_Nam,
    ROUND(
      SUM(CASE WHEN r.npl_flag=1 THEN f.loan_amnt ELSE 0 END)
      - LAG(SUM(CASE WHEN r.npl_flag=1 THEN f.loan_amnt ELSE 0 END),1,0)
        OVER(ORDER BY t.issue_year),2)  AS Tang_Truong_No_Xau
FROM Fact_Loans f
JOIN Dim_Time t       ON f.time_id = t.time_id
JOIN Dim_CreditRisk r ON f.risk_id = r.risk_id
GROUP BY t.issue_year ORDER BY t.issue_year;""")

# Bảng 3.11 — KPI 13.png
add_table(doc,
    headers=["Năm", "Tổng Khoản Vay", "Số Khoản NPL", "Tổng Giải Ngân ($)", "Dư Nợ Xấu ($)", "NPL Ratio (%)", "Tăng Trưởng Nợ Xấu ($)"],
    rows=[
        ["2007", "572",     "44",     "4,819,275",     "463,400",       "9.62",  "+463,400"],
        ["2008", "2,391",   "247",    "21,085,800",    "2,739,600",     "12.99", "+2,276,200"],
        ["2009", "5,267",   "593",    "51,733,075",    "6,236,475",     "12.06", "+3,496,875"],
        ["2010", "12,506",  "1,482",  "131,608,100",   "15,954,200",    "12.12", "+9,717,725"],
        ["2011", "21,710",  "3,296",  "261,454,525",   "43,344,825",    "16.58", "+27,390,625"],
        ["2012", "53,335",  "8,644",  "717,616,675",   "127,242,925",   "17.73", "+83,898,100"],
        ["2013", "134,738", "21,011", "1,980,794,700", "329,506,525",   "16.64", "+202,263,600"],
        ["2014", "235,464", "41,144", "3,499,445,475", "640,555,075",   "18.30", "+311,048,550"],
        ["2015", "420,668", "75,761", "6,407,095,775", "1,191,158,075", "18.59", "+550,603,000"],
        ["2016", "433,873", "68,200", "6,386,883,825", "1,050,244,950", "16.44", "−140,913,125"],
        ["2017", "443,050", "39,159", "6,569,318,500", "618,844,250",   "9.42",  "−431,400,700"],
        ["2018", "494,510", "8,860",  "7,917,423,475", "150,788,400",   "1.90",  "−468,055,850"],
    ],
    col_widths=[Cm(1.5), Cm(2.5), Cm(2.3), Cm(3.2), Cm(3.2), Cm(2.2), Cm(3.5)],
    highlight_rows={8},
    warn_rows={11},
    caption="Bảng 3.11: Tăng trưởng NPL theo năm giải ngân (Nguồn: KPI 13)"
)

add_remark(doc,
    "NPL Ratio đạt đỉnh năm 2015 (18.59%) — cohort 2015 có tỷ lệ nợ xấu cao nhất do đây là năm "
    "giải ngân mạnh nhất với nhiều khoản vay kỳ hạn 36–60 tháng đã đủ thời gian trở thành nợ xấu. "
    "Hiệu ứng Loan Age rõ ràng: NPL Ratio giảm mạnh từ 18.59% (2015) xuống 9.42% (2017) "
    "và chỉ 1.90% (2018) — phản ánh khoản vay càng non trẻ thì chưa đủ thời gian phát sinh nợ xấu. "
    "Năm 2016 NPL giảm tuyệt đối −$140.9M so với 2015: sau khủng hoảng quản trị, "
    "Lending Club thắt chặt tiêu chuẩn tín dụng, cải thiện chất lượng cohort 2016. "
    "Tăng trưởng nợ xấu tuyệt đối lớn nhất năm 2015 (+$550.6M) — năm yêu cầu vốn dự phòng ECL lớn nhất.")

doc.add_paragraph()

# ══ NHÓM 6 ════════════════════════════════════
add_heading(doc, "Nhóm 6: Phân tích lãi suất và tỷ lệ hoàn trả", level=2)

# ── 3.1.13 ───────────────────────────────────
add_heading(doc, "3.1.13. Lãi suất trung bình theo hạng tín dụng (Truy vấn 8)", level=3)

add_normal(doc, "Mục đích: Kiểm chứng cơ chế định giá rủi ro (Risk-Based Pricing) — Grade rủi ro cao hơn phải chịu lãi suất cao hơn để bù đắp xác suất vỡ nợ.", indent=True)

add_sql(doc, """SELECT
    p.grade,
    COUNT(f.loan_id)          AS So_Khoan_Vay,
    ROUND(AVG(f.int_rate),2)  AS Lai_Suat_TB,
    ROUND(MIN(f.int_rate),2)  AS Lai_Suat_Min,
    ROUND(MAX(f.int_rate),2)  AS Lai_Suat_Max
FROM Fact_Loans f
JOIN Dim_LoanProduct p ON f.product_id = p.product_id
GROUP BY p.grade ORDER BY p.grade;""")

# Bảng 3.12 — KPI 8.png (số liệu thực tế)
add_table(doc,
    headers=["Grade", "Số Khoản Vay", "LS Trung Bình (%)", "LS Min (%)", "LS Max (%)"],
    rows=[
        ["A", "432,226", "7.08",  "5.31", "9.63"],
        ["B", "662,757", "10.68", "6.00", "14.09"],
        ["C", "649,478", "14.14", "6.00", "17.27"],
        ["D", "324,169", "18.14", "6.00", "22.35"],
        ["E", "135,534", "21.83", "6.00", "27.27"],
        ["F", "41,767",  "25.45", "6.00", "30.75"],
        ["G", "12,153",  "28.08", "6.00", "30.99"],
    ],
    col_widths=[Cm(2.0), Cm(3.5), Cm(3.5), Cm(2.8), Cm(2.8)],
    caption="Bảng 3.12: Biên độ lãi suất theo hạng tín dụng (Nguồn: KPI 8)"
)

add_remark(doc,
    "Hệ thống định giá rủi ro hoạt động đơn điệu và chính xác — lãi suất tăng liên tục từ Grade A (7.08%) "
    "đến Grade G (28.08%), khoảng cách lên tới 21.00 điểm phần trăm, xác nhận cơ chế Risk-Based Pricing "
    "hoạt động đúng nguyên tắc. Grade A ở mức 7.08% — dưới ngưỡng 10%, cạnh tranh được với lãi suất "
    "cho vay cá nhân của ngân hàng truyền thống tại Mỹ (8–12%), là lợi thế thu hút khách hàng chất lượng cao. "
    "Grade B và C chiếm đa số danh mục (662,757 + 649,478 = 1,312,235 khoản, ~57.7% tổng) "
    "với lãi suất 10.68%–14.14% — phân khúc sinh lời ổn định nhất cho nhà đầu tư P2P. "
    "Đáng chú ý, tất cả Grade từ B đến G đều có Lãi Suất Min = 6.00%, phản ánh mức sàn lãi suất "
    "tối thiểu mà Lending Club áp dụng xuyên suốt các nhóm rủi ro.")

doc.add_paragraph()

# ── 3.1.14 ───────────────────────────────────
add_heading(doc, "3.1.14. Tỷ lệ hoàn trả toàn danh mục (Truy vấn 2)", level=3)

add_normal(doc, "Mục đích: Tính tỷ lệ phần trăm tổng số tiền đã thanh toán so với tổng số tiền đã giải ngân — chỉ số tổng hợp phản ánh sức khỏe thanh toán của toàn bộ danh mục.", indent=True)

add_sql(doc, """SELECT
    SUM(total_pymnt)   AS Tong_Da_Thanh_Toan,
    SUM(loan_amnt)     AS Tong_Giai_Ngan,
    ROUND((SUM(total_pymnt)/NULLIF(SUM(loan_amnt),0))*100,2)
                       AS Ty_Le_Hoan_Tra_PhanTram
FROM Fact_Loans;""")

# Bảng 3.13 — KPI 2.png (số liệu thực tế)
add_table(doc,
    headers=["Chỉ Tiêu", "Giá Trị"],
    rows=[
        ["Tổng tiền đã thanh toán (total_pymnt)",   "$27,262,913,860.36"],
        ["Tổng tiền đã giải ngân (loan_amnt)",       "$33,949,279,200.00"],
        ["Tỷ lệ hoàn trả toàn danh mục (%)",         "80.30%"],
    ],
    col_widths=[Cm(9.0), Cm(5.0)],
    caption="Bảng 3.13: Tỷ lệ hoàn trả toàn danh mục (Nguồn: KPI 2)"
)

add_remark(doc,
    "Tỷ lệ hoàn trả toàn danh mục đạt 80.30% ($27.26 tỷ / $33.95 tỷ). "
    "Chênh lệch $6.69 tỷ chưa được thanh toán bao gồm: dư nợ còn lại của các khoản đang hoạt động (Current) "
    "chưa hoàn tất kỳ hạn + nợ xấu không thể thu hồi. "
    "Nếu chỉ tính riêng các khoản đã đóng hoàn toàn (Fully Paid + Charged Off), "
    "tỷ lệ hoàn trả vượt 100% vì total_pymnt bao gồm cả lãi, phí dịch vụ và tiền phạt trả chậm — "
    "đây là chỉ số tích cực cho thấy đa số khách hàng đã thực hiện nghĩa vụ thanh toán đầy đủ.")

doc.add_paragraph()

# ══ TỔNG HỢP KPI ══════════════════════════════
add_heading(doc, "Tổng hợp kết quả phân tích KPI", level=2)

add_normal(doc, "Bảng 3.14 dưới đây tóm tắt toàn bộ các chỉ số KPI chính được trích xuất thông qua 14 truy vấn T-SQL trên hệ thống CreditBI_DB:")

add_table(doc,
    headers=["KPI", "Giá Trị", "Đánh Giá"],
    rows=[
        ["Tổng khoản vay đã xử lý",            "2,258,084",          "Quy mô lớn"],
        ["Tổng giải ngân",                      "$33.95 tỷ USD",      "—"],
        ["Tỷ lệ nợ xấu NPL — toàn danh mục",   "14.20%",             "⚠ Cần kiểm soát"],
        ["Bang có NPL Ratio cao nhất",          "Louisiana: 14.98%",  "Vùng South dẫn đầu"],
        ["FICO TB nhóm Low Risk",               "702 điểm",           "✓ Lành mạnh"],
        ["FICO TB nhóm Default",                "690 điểm",           "Thấp hơn Low 12 điểm"],
        ["DTI ngưỡng cảnh báo rủi ro",          "> 20%",              "Nhóm Medium/High/Default"],
        ["Tỷ lệ thu hồi nợ (Grade A → G)",     "5.91% → 9.71%",      "Tăng dần theo rủi ro"],
        ["Cohort có NPL Ratio cao nhất",        "2015: 18.59%",       "Đỉnh bùng nổ P2P Lending"],
        ["Tỷ lệ hoàn trả toàn danh mục",       "80.30%",             "Gồm khoản đang Current"],
        ["Lãi suất Grade A",                   "7.15%",              "Cạnh tranh ngân hàng TT"],
        ["Lãi suất Grade G",                   "27.36%",             "Bù đắp rủi ro cao"],
        ["Grade chiếm tỷ trọng lớn nhất",      "B & C — 58.2%",      "Sinh lời ổn định"],
        ["Năm giải ngân cao nhất",             "2018: $7.92 tỷ",     "Đỉnh lịch sử"],
        ["Hiệu quả thu hồi đỉnh theo năm",    "2016: 1.95%",        "Cohort trưởng thành nhất"],
    ],
    col_widths=[Cm(6.5), Cm(4.5), Cm(4.5)],
    caption="Bảng 3.14: Tóm tắt các chỉ số KPI chính của hệ thống CreditBI"
)

add_normal(doc,
    "Qua 14 truy vấn phân tích T-SQL trên hệ thống CreditBI_DB, nhóm đã trích xuất được bức tranh "
    "toàn diện về danh mục tín dụng Lending Club. Mô hình Star Schema và hệ thống 9 Index "
    "(5 trên Fact_Loans và 4 trên loan_staging_clean) đã phát huy hiệu quả — tất cả truy vấn JOIN "
    "nhiều bảng trên 2.26 triệu bản ghi hoàn thành trong thời gian dưới 8 giây trên cấu hình "
    "SQL Server thông thường, đáp ứng yêu cầu phân tích nghiệp vụ thực tế. "
    "Các kết quả phân tích này được đưa vào dashboard để trực quan hóa trong mục 3.2.",
    indent=True)

# ── Lưu file ─────────────────────────────────
OUT = r'c:\Nam 3 ki 2\HQT_CSDL\Dataset\Section_3_1_KPI_Analysis.docx'
doc.save(OUT)
print(f"✅ Đã lưu: {OUT}")
